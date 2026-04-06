from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


TITLE = "Real-Time Emotion-Based Image Matching Using a Webcam and ONNX FER+"

ABSTRACT_CONTENT = (
    "This project presents a real-time computer vision system that recognizes a user's "
    "facial emotion from a webcam stream and matches it with the most emotionally similar "
    "image from a local dataset. The purpose of the system is to create an interactive "
    "human-computer interface in which visual content dynamically adapts to the user's "
    "emotional state. The proposed solution is implemented in Python and combines "
    "MediaPipe face detection, ONNX FER+ emotion recognition, and cosine-similarity-based "
    "matching of emotion vectors. Since the original FER+ model predicts eight classes, "
    "its output is mapped to the standard seven-emotion representation: angry, disgust, "
    "fear, happy, neutral, sad, and surprise. This approach preserves compatibility with "
    "common emotion-labeling schemes while avoiding TensorFlow dependency issues in the "
    "target Python 3.14 environment. To improve performance, dataset embeddings are "
    "indexed in advance and stored in a cache, while computationally expensive inference "
    "is not executed on every frame, which helps maintain smooth real-time interaction. "
    "The system also includes practical interface features such as smooth preview "
    "transitions, optional sound alerts for strong matches, and logging of recognized "
    "emotions. The developed prototype demonstrates that emotion-based image matching can "
    "be implemented with accessible tools and modest computational requirements, making it "
    "a useful example of applied affective computing and real-time computer vision."
)

ACKNOWLEDGMENTS = (
    "The author expresses gratitude to Affiliation with the university \"AIU\" "
    "(Kazakhstan) for academic support and for providing an environment for the "
    "development of this project."
)

REFERENCES = [
    "Mollahosseini A., Hasani B., Mahoor M. H. AffectNet: A Database for Facial Expression, Valence, and Arousal Computing in the Wild. IEEE Transactions on Affective Computing, 10(1), 18-31, 2019.",
    "Lugaresi C., Tang J., Nash H., et al. MediaPipe: A Framework for Building Perception Pipelines. arXiv:1906.08172, 2019.",
    "ONNX Model Zoo. Emotion FER+ Model. GitHub repository: https://github.com/onnx/models",
    "Bradski G. The OpenCV Library. Dr. Dobb's Journal of Software Tools, 2000.",
]

KEYWORDS = (
    "facial emotion recognition; affective computing; real-time computer vision; "
    "MediaPipe; ONNX FER+; cosine similarity; webcam interaction; image matching"
)

TRACK = "Recommended: Artificial Intelligence / Computer Vision / Machine Learning"
PRESENTATION_TYPE = "Recommended: Oral presentation"
HOW_TO_PRESENT = "Recommended: In-person presentation"


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

add_heading(doc, "Abstract title*")
add_body(doc, TITLE)

add_heading(doc, "Abstract content*")
add_body(doc, ABSTRACT_CONTENT)

add_heading(doc, "Acknowledgments")
add_body(doc, ACKNOWLEDGMENTS)

add_heading(doc, "References")
for ref in REFERENCES:
    add_body(doc, ref)

add_heading(doc, "Abstract keywords")
add_body(doc, KEYWORDS)

add_heading(doc, "Conference track*")
add_body(doc, TRACK)

add_heading(doc, "Presentation type*")
add_body(doc, PRESENTATION_TYPE)

add_heading(doc, "How do you want to present?*")
add_body(doc, HOW_TO_PRESENT)

output_path = "submission_content.docx"
doc.save(output_path)
print(output_path)
