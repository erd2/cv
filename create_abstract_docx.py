"""Скрипт для создания abstract.docx из текста абстракта."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Настройка стилей
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.space_before = Pt(0)
paragraph_format.line_spacing = 1.5

# --- TITLE ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Real-Time Emotion-Based Image Matching Using a Webcam and ONNX FER+')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

# --- ABSTRACT BODY ---
abstract_paragraphs = [
    (
        "This project presents a real-time computer vision system that recognizes a user's "
        "facial emotion from a webcam stream and matches it with the most emotionally similar "
        "image from a local dataset. The main goal of the system is to create an interactive "
        "human-computer interface in which visual content dynamically adapts to the emotional "
        "state of the user. Such an approach can be applied in entertainment software, "
        "interactive installations, personalized media systems, and educational demonstrations "
        "of affective computing."
    ),
    (
        "The proposed solution is implemented in Python and combines several practical "
        "components into a single lightweight pipeline. Face detection is performed with "
        "MediaPipe, which provides stable and fast localization of faces in live video. "
        "Emotion recognition is based on the ONNX FER+ model. Since the original FER+ network "
        "predicts eight classes, its output is mapped to the common seven-emotion representation: "
        "angry, disgust, fear, happy, neutral, sad, and surprise. This design preserves "
        "compatibility with standard emotion-labeling schemes while avoiding the dependency on "
        "TensorFlow, which is inconvenient for the target Python 3.14 environment. For each "
        "detected face, the system forms an emotion-probability vector and compares it with "
        "vectors precomputed for images in the local dataset."
    ),
    (
        "To improve performance, dataset embeddings are indexed in advance and stored in a "
        "cache. During runtime, the application processes the webcam stream in real time, limits "
        "expensive model calls to maintain responsiveness, and performs similarity search using "
        "cosine similarity. The system also includes practical interface features such as smooth "
        "preview transitions, optional sound notification for strong matches, logging of "
        "recognized emotions, and support for manual emotion labels in the dataset metadata."
    ),
    (
        "The developed prototype demonstrates that emotion-based matching can be implemented "
        "with accessible tools and modest computational requirements while preserving stable "
        "interactive behavior. The project shows how affective analysis, efficient indexing, "
        "and real-time visualization can be integrated into a compact applied system, making "
        "emotion-aware interaction more understandable and accessible for further research and "
        "development."
    ),
]

for text in abstract_paragraphs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(36)  # ~1.27 cm абзацный отступ
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

# --- KEYWORDS ---
keywords_label = doc.add_paragraph()
keywords_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = keywords_label.add_run('Keywords: ')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = keywords_label.add_run(
    'facial emotion recognition, affective computing, real-time computer vision, '
    'MediaPipe, ONNX FER+, cosine similarity, webcam-based interaction, emotion-based image matching'
)
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
keywords_label.paragraph_format.space_before = Pt(12)

# Сохранение
output_path = 'abstract_en.docx'
doc.save(output_path)
print(f"Файл создан: {output_path}")
